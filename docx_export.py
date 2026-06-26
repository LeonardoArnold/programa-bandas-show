"""
Geração da ficha em Word (.docx) no formato do Programa Bandas Show.

Usa python-docx (única dependência além do FastAPI/uvicorn).
A numeração das bandas é calculada na hora: só blocos do tipo "banda" recebem número.
"""
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paleta
VERMELHO = RGBColor(0xC0, 0x00, 0x00)
PRETO = RGBColor(0x00, 0x00, 0x00)
FONTE = "Arial"
TAM_TITULO = 14
TAM_LINHA = 12


def _run(par, texto, *, bold=True, cor=PRETO, tam=TAM_LINHA, highlight=False):
    r = par.add_run(texto)
    r.bold = bold
    r.font.name = FONTE
    r.font.size = Pt(tam)
    r.font.color.rgb = cor
    if highlight:
        # 7 = amarelo (WD_COLOR_INDEX.YELLOW)
        r.font.highlight_color = 7
    return r


def gerar_docx(ficha: dict) -> bytes:
    """
    ficha = {
        "numero": "1110",
        "data_exibicao": "28/06/2026",   # já formatado dd/mm/aaaa
        "data_gravacao": "20/06/2026",
        "blocos": [
            {"tipo": "banda", "nome": "BANDA MARIAH", "musica": "GARÇOM",
             "lancamento": False, "feat": ""},
            {"tipo": "comercial"},
            {"tipo": "texto", "conteudo": "FRIGORIFICO FRIGOBENDO"},
            ...
        ]
    }
    """
    doc = Document()

    # Margens e fonte padrão
    style = doc.styles["Normal"]
    style.font.name = FONTE
    style.font.size = Pt(TAM_LINHA)

    # --- Cabeçalho ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(p, f"PROGRAMA N°({ficha.get('numero', '')})",
         bold=True, cor=VERMELHO, tam=TAM_TITULO, highlight=True)

    # Linha opcional com datas (deixei discreta; some se vazias)
    if ficha.get("data_exibicao") or ficha.get("data_gravacao"):
        pd = doc.add_paragraph()
        partes = []
        if ficha.get("data_exibicao"):
            partes.append(f"Exibição: {ficha['data_exibicao']}")
        if ficha.get("data_gravacao"):
            partes.append(f"Gravação: {ficha['data_gravacao']}")
        _run(pd, "   |   ".join(partes), bold=False, cor=PRETO, tam=10)

    doc.add_paragraph()  # linha em branco

    # --- Blocos ---
    n_banda = 0
    for bloco in ficha.get("blocos", []):
        tipo = bloco.get("tipo")
        p = doc.add_paragraph()

        if tipo == "banda":
            n_banda += 1
            nome = (bloco.get("nome") or "").upper().strip()
            musica = (bloco.get("musica") or "").upper().strip()
            texto = f"{n_banda}-{nome}"
            if musica:
                texto += f" ({musica})"
            _run(p, texto, bold=True, cor=PRETO)
            feat = (bloco.get("feat") or "").upper().strip()
            if feat:
                _run(p, f" FT. {feat}", bold=True, cor=PRETO)
            if bloco.get("lancamento"):
                _run(p, " LANÇAMENTO", bold=True, cor=VERMELHO)

        elif tipo == "comercial":
            _run(p, "COMERCIAL", bold=True, cor=PRETO)

        elif tipo == "redes_sociais":
            txt = (bloco.get("conteudo") or "REDES SOCIAIS").upper().strip()
            _run(p, txt, bold=True, cor=PRETO)

        elif tipo == "abracos":
            txt = (bloco.get("conteudo") or "ABRAÇOS").upper().strip()
            _run(p, txt, bold=True, cor=PRETO)

        else:  # texto, contato, ou qualquer bloco livre
            txt = (bloco.get("conteudo") or "").upper().strip()
            _run(p, txt, bold=True, cor=PRETO)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


if __name__ == "__main__":
    # Teste rápido reproduzindo a ficha 1109 da foto
    exemplo = {
        "numero": "1109",
        "data_exibicao": "21/06/2026",
        "data_gravacao": "13/06/2026",
        "blocos": [
            {"tipo": "banda", "nome": "Banda Mariah", "musica": "Garçom"},
            {"tipo": "banda", "nome": "Integração", "musica": "Judiaria"},
            {"tipo": "comercial"},
            {"tipo": "banda", "nome": "Os Atuais", "musica": "Voarei"},
            {"tipo": "texto", "conteudo": "Contato equipe de marketing (whatsapp) (45) 991160918"},
            {"tipo": "texto", "conteudo": "Frigorifico Frigobendo"},
            {"tipo": "banda", "nome": "Teixerinha", "musica": "Querência Amada"},
            {"tipo": "texto", "conteudo": "Bailão e bandinha com Junior Ghesla"},
            {"tipo": "texto", "conteudo": "Erva Mate Flor Serrana"},
            {"tipo": "comercial"},
            {"tipo": "banda", "nome": "Banda Milennium", "musica": "Fogo e Gasolina"},
            {"tipo": "banda", "nome": "Banda Elyte", "musica": "Tome Gelo"},
            {"tipo": "abracos", "conteudo": "Abraços"},
            {"tipo": "redes_sociais", "conteudo": "Redes Sociais (Rogerio Arnold)"},
            {"tipo": "banda", "nome": "Musical San Francisco", "musica": "Dor Bandida", "lancamento": True},
            {"tipo": "comercial"},
            {"tipo": "banda", "nome": "Musical Calmon", "musica": "Mesa do Canto"},
            {"tipo": "banda", "nome": "Gaucho da Fronteira", "musica": "Adeus Mariana", "feat": "Leonardo"},
            {"tipo": "texto", "conteudo": "Martinho Francisco"},
        ],
    }
    with open("exemplo_1109.docx", "wb") as f:
        f.write(gerar_docx(exemplo))
    print("exemplo_1109.docx gerado")
